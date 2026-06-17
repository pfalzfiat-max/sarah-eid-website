"""
Erstellt ANLEITUNG-STUDIO.docx mit professionellem Design via python-docx.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Farben ──────────────────────────────────────────────────
GOLD        = RGBColor(0xC9, 0xA8, 0x4C)
DARK        = RGBColor(0x1A, 0x1A, 0x2E)
BODY        = RGBColor(0x2D, 0x2D, 0x3A)
MUTED       = RGBColor(0x5A, 0x5A, 0x70)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG    = RGBColor(0xF8, 0xF6, 0xF1)
BORDER      = RGBColor(0xE0, 0xD8, 0xC0)
WARN_BG     = RGBColor(0xFF, 0xF8, 0xE8)
WARN_BORDER = RGBColor(0xC9, 0xA8, 0x4C)

# ─── Hilfsfunktionen ─────────────────────────────────────────

def set_cell_bg(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = str(color)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 6)))
            el.set(qn('w:color'), val.get('color', 'C9A84C'))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def set_para_border_left(para, color='C9A84C', sz=18):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), str(sz))
    left.set(qn('w:space'), '12')
    left.set(qn('w:color'), color)
    pBdr.append(left)
    pPr.append(pBdr)

def set_para_shading(para, fill_color: RGBColor):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    hex_color = str(fill_color)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

def add_run(para, text, bold=False, italic=False, color=None, size=None):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return run

def para_space(para, before=0, after=0, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        from docx.shared import Pt as DPt
        pf.line_spacing = DPt(line)

def set_font(run, name='Calibri'):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)

# ─── Dokument erstellen ───────────────────────────────────────

doc = Document()

# Seitenränder
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# Standard-Schrift für alle Styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = BODY

# ─── HEADER ──────────────────────────────────────────────────

# Goldene Titelbox
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
para_space(title_para, before=0, after=4, line=16)
set_para_shading(title_para, DARK)
r = title_para.add_run('  Website-Anleitung')
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = GOLD
r.font.name = 'Calibri'

subtitle_para = doc.add_paragraph()
subtitle_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
para_space(subtitle_para, before=0, after=16, line=14)
set_para_shading(subtitle_para, DARK)
r2 = subtitle_para.add_run('  saraheid.de – Inhalte selbst bearbeiten')
r2.font.size = Pt(12)
r2.font.color.rgb = RGBColor(0xC0, 0xB8, 0xA0)
r2.font.name = 'Calibri'

# Goldene Trennlinie
line_para = doc.add_paragraph()
para_space(line_para, before=0, after=20)
line_run = line_para.add_run('─' * 80)
line_run.font.color.rgb = GOLD
line_run.font.size = Pt(8)

# ─── ABSCHNITT: EINLOGGEN ────────────────────────────────────

def section_heading(doc, text):
    p = doc.add_paragraph()
    para_space(p, before=18, after=6)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = GOLD
    r.font.name = 'Calibri'
    r.font.letter_spacing = Pt(1)
    # Unterstrich
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), 'C9A84C')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def body_para(doc, text='', indent=False):
    p = doc.add_paragraph()
    para_space(p, before=2, after=6, line=14)
    if indent:
        p.paragraph_format.left_indent = Cm(0.6)
    if text:
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.color.rgb = BODY
        r.font.name = 'Calibri'
    return p

def info_box(doc, text):
    """Gold-Randkasten links für Hinweise."""
    p = doc.add_paragraph()
    para_space(p, before=8, after=8, line=14)
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.4)
    set_para_border_left(p, color='C9A84C', sz=18)
    set_para_shading(p, LIGHT_BG)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = BODY
    r.font.name = 'Calibri'
    return p

def warn_box(doc, text):
    """Warnkasten."""
    p = doc.add_paragraph()
    para_space(p, before=8, after=8, line=14)
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.4)
    set_para_border_left(p, color='C94C4C', sz=18)
    set_para_shading(p, RGBColor(0xFF, 0xF0, 0xF0))
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(0x6B, 0x20, 0x20)
    r.font.name = 'Calibri'
    return p

def step_item(doc, number, text):
    p = doc.add_paragraph()
    para_space(p, before=3, after=3, line=14)
    p.paragraph_format.left_indent = Cm(0.6)
    r_num = p.add_run(f'{number}   ')
    r_num.bold = True
    r_num.font.size = Pt(11)
    r_num.font.color.rgb = GOLD
    r_num.font.name = 'Calibri'
    r_text = p.add_run(text)
    r_text.font.size = Pt(11)
    r_text.font.color.rgb = BODY
    r_text.font.name = 'Calibri'
    return p

def bullet_item(doc, text, bold_part=None):
    p = doc.add_paragraph()
    para_space(p, before=3, after=3, line=14)
    p.paragraph_format.left_indent = Cm(0.6)
    r_bullet = p.add_run('·   ')
    r_bullet.font.size = Pt(11)
    r_bullet.font.color.rgb = GOLD
    r_bullet.font.name = 'Calibri'
    if bold_part:
        r_b = p.add_run(bold_part)
        r_b.bold = True
        r_b.font.size = Pt(11)
        r_b.font.color.rgb = BODY
        r_b.font.name = 'Calibri'
        r_rest = p.add_run(text)
        r_rest.font.size = Pt(11)
        r_rest.font.color.rgb = BODY
        r_rest.font.name = 'Calibri'
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.color.rgb = BODY
        r.font.name = 'Calibri'
    return p


# ══════════════════════════════════════════════════════════════
# 1. EINLOGGEN
# ══════════════════════════════════════════════════════════════

section_heading(doc, '1   Einloggen')

p = body_para(doc)
p.add_run('Adresse im Browser öffnen: ').font.size = Pt(11)
p.runs[0].font.color.rgb = BODY
p.runs[0].font.name = 'Calibri'
r_url = p.add_run('saraheid.de/studio')
r_url.bold = True
r_url.font.size = Pt(11)
r_url.font.color.rgb = GOLD
r_url.font.name = 'Calibri'

p2 = body_para(doc)
r2a = p2.add_run('Du hast eine Einladungs-E-Mail von Sanity an ')
r2a.font.size = Pt(11)
r2a.font.color.rgb = BODY
r2a.font.name = 'Calibri'
r2b = p2.add_run('mail@saraheid.de')
r2b.bold = True
r2b.font.size = Pt(11)
r2b.font.color.rgb = BODY
r2b.font.name = 'Calibri'
r2c = p2.add_run(' erhalten. Bitte zuerst den Link in dieser E-Mail bestätigen – danach ist der Login freigeschaltet.')
r2c.font.size = Pt(11)
r2c.font.color.rgb = BODY
r2c.font.name = 'Calibri'

p3 = body_para(doc)
r3a = p3.add_run('Beim Login kannst du ein Google-Konto mit der Adresse ')
r3a.font.size = Pt(11)
r3a.font.color.rgb = BODY
r3a.font.name = 'Calibri'
r3b = p3.add_run('mail@saraheid.de')
r3b.bold = True
r3b.font.size = Pt(11)
r3b.font.color.rgb = BODY
r3b.font.name = 'Calibri'
r3c = p3.add_run(' verwenden oder ein eigenes Passwort setzen.')
r3c.font.size = Pt(11)
r3c.font.color.rgb = BODY
r3c.font.name = 'Calibri'

# ══════════════════════════════════════════════════════════════
# 2. WAS DU BEARBEITEN KANNST
# ══════════════════════════════════════════════════════════════

section_heading(doc, '2   Was du bearbeiten kannst')

p_intro = body_para(doc, 'Im Studio siehst du links eine Übersicht aller Bereiche:')

# Tabelle
table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.LEFT
table.style = 'Table Grid'

# Header-Zeile
hdr = table.rows[0].cells
set_cell_bg(hdr[0], DARK)
set_cell_bg(hdr[1], DARK)
for cell, text in [(hdr[0], 'Bereich'), (hdr[1], 'Was du dort änderst')]:
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = GOLD
    r.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    cell.width = Inches(2.2) if text == 'Bereich' else Inches(4.0)

rows_data = [
    ('Einstellungen', 'Name, Telefon, E-Mail, Social-Media-Links, SEO-Texte, Footer-Text, alle Bilder'),
    ('Startseite', 'Hero-Tagline, Untertitel, Beschreibungstext, Über-mich-Texte, Politik-Texte'),
    ('Leistungen', 'Die Leistungskarten (Titel, Beschreibung, Reihenfolge)'),
    ('Events', 'Aktuelle Veranstaltungen mit Bild, Beschreibung und Kategorie'),
    ('Referenzen', 'Kundenstimmen (Text oder Vimeo-Video)'),
    ('FAQ', 'Häufige Fragen und Antworten'),
    ('Impressum', 'Adresse, Telefon, E-Mail, Steuernummer'),
    ('Datenschutz', 'Datenschutzerklärung (nur nach rechtlicher Prüfung ändern)'),
]

for i, (bereich, was) in enumerate(rows_data):
    row = table.add_row()
    bg = LIGHT_BG if i % 2 == 0 else WHITE
    for cell in row.cells:
        set_cell_bg(cell, bg)
    c0, c1 = row.cells
    p0 = c0.paragraphs[0]
    p0.paragraph_format.space_before = Pt(4)
    p0.paragraph_format.space_after = Pt(4)
    r0 = p0.add_run(bereich)
    r0.bold = True
    r0.font.size = Pt(10.5)
    r0.font.color.rgb = DARK
    r0.font.name = 'Calibri'

    p1 = c1.paragraphs[0]
    p1.paragraph_format.space_before = Pt(4)
    p1.paragraph_format.space_after = Pt(4)
    r1 = p1.add_run(was)
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = BODY
    r1.font.name = 'Calibri'

doc.add_paragraph()  # Abstand nach Tabelle

# ══════════════════════════════════════════════════════════════
# 3. TEXTE ÄNDERN
# ══════════════════════════════════════════════════════════════

section_heading(doc, '3   Texte ändern')

for n, text in [
    ('1', 'Im Studio links den passenden Bereich anklicken'),
    ('2', 'Das Dokument öffnet sich rechts'),
    ('3', 'Einfach in das Textfeld klicken und den Text bearbeiten'),
    ('4', 'Oben rechts auf Veröffentlichen klicken'),
]:
    step_item(doc, n, text)

info_box(doc,
    '⏱  Änderungen sind in der Regel nach etwa einer Minute live auf der Website sichtbar. '
    'Seite einfach neu laden (Strg+Shift+R bzw. Cmd+Shift+R).')

# ══════════════════════════════════════════════════════════════
# 4. BILDER HOCHLADEN
# ══════════════════════════════════════════════════════════════

section_heading(doc, '4   Bilder hochladen')

p_b1 = body_para(doc, 'Alle Hauptbilder der Website findest du unter Einstellungen:')

for bold, rest in [
    ('Hero – Hintergrundbild', ' → atmosphärisches Bild hinter dem Titel'),
    ('Hero – Portrait-Foto', ' → dein Portrait rechts im Einstiegsbereich'),
    ('Über mich – Bild links oben', ' → erstes Foto in der Über-mich-Sektion'),
    ('Über mich – Bild rechts unten', ' → zweites Foto in der Über-mich-Sektion'),
]:
    bullet_item(doc, rest, bold_part=bold)

p_spacer = body_para(doc, 'So lädst du ein Bild hoch:')
para_space(p_spacer, before=10, after=4)

for n, text in [
    ('1', 'Einstellungen öffnen'),
    ('2', 'Nach unten scrollen bis zum gewünschten Bildfeld'),
    ('3', 'Auf Upload klicken und das Bild vom Computer auswählen – oder Bild direkt ins Feld ziehen'),
    ('4', 'Fokuspunkt (Hotspot) setzen: bestimmt den Bildausschnitt bei verschiedenen Bildschirmgrößen'),
    ('5', 'Alt-Text eintragen – kurze Bildbeschreibung für Google und Barrierefreiheit\n        (Beispiel: „Sarah Eid moderiert auf der Bühne")'),
    ('6', 'Oben rechts Veröffentlichen klicken'),
]:
    step_item(doc, n, text)

info_box(doc,
    '📸  Die Bilder auf der Website sind derzeit noch Platzhalter. '
    'Sobald du ein Bild im Studio hochlädst und veröffentlichst, ersetzt es automatisch den Platzhalter – '
    'du musst nichts weiter tun.')

# ══════════════════════════════════════════════════════════════
# 5. NEUE INHALTE HINZUFÜGEN
# ══════════════════════════════════════════════════════════════

section_heading(doc, '5   Neue Leistung oder Event hinzufügen')

for n, text in [
    ('1', 'Im linken Menü auf Leistungen oder Events klicken'),
    ('2', 'Oben rechts auf das + (Neu erstellen) klicken'),
    ('3', 'Felder ausfüllen und Veröffentlichen klicken'),
    ('4', 'Reihenfolge: niedrige Zahl erscheint zuerst'),
]:
    step_item(doc, n, text)

# ══════════════════════════════════════════════════════════════
# 6. HÄUFIGE FRAGEN
# ══════════════════════════════════════════════════════════════

section_heading(doc, '6   Häufige Fragen')

for frage, antwort in [
    (
        'Ich habe etwas geändert – warum sehe ich es noch nicht?',
        'Warte ca. eine Minute und lade die Website neu. Die Website aktualisiert sich automatisch, aber nicht sofort.'
    ),
    (
        'Ich habe etwas aus Versehen falsch geändert.',
        'Im Studio gibt es eine Versionshistorie. Klicke oben im Dokument auf den Pfeil neben „Veröffentlicht" – dort kannst du ältere Versionen wiederherstellen.'
    ),
    (
        'Ein Bild wird nicht angezeigt.',
        'Prüfe, ob du nach dem Hochladen auf Veröffentlichen geklickt hast. Unveröffentlichte Änderungen sind noch nicht live.'
    ),
]:
    p_q = body_para(doc)
    p_q.paragraph_format.left_indent = Cm(0)
    r_q = p_q.add_run(frage)
    r_q.bold = True
    r_q.font.size = Pt(11)
    r_q.font.color.rgb = DARK
    r_q.font.name = 'Calibri'
    para_space(p_q, before=10, after=2)

    p_a = body_para(doc, antwort)
    p_a.paragraph_format.left_indent = Cm(0.6)
    para_space(p_a, before=0, after=8)

# ══════════════════════════════════════════════════════════════
# 7. HINWEIS IMPRESSUM / DATENSCHUTZ
# ══════════════════════════════════════════════════════════════

section_heading(doc, '7   Impressum und Datenschutz')

warn_box(doc,
    '⚠  Impressum und Datenschutzerklärung sind im Studio editierbar – bitte ändere diese Texte nur, '
    'wenn du die Änderung vorher rechtlich geprüft oder von einem Anwalt freigeben lassen hast. '
    'Fehler in diesen Bereichen können rechtliche Konsequenzen haben.')

# ══════════════════════════════════════════════════════════════
# 8. SUPPORT
# ══════════════════════════════════════════════════════════════

section_heading(doc, '8   Support')

p_sup = body_para(doc, 'Bei technischen Problemen oder Fragen bitte direkt Michael Weiss kontaktieren.')

# Abschlusslinie
doc.add_paragraph()
line_para2 = doc.add_paragraph()
para_space(line_para2, before=16, after=4)
line_run2 = line_para2.add_run('─' * 80)
line_run2.font.color.rgb = RGBColor(0xD0, 0xC8, 0xA8)
line_run2.font.size = Pt(8)

p_footer = doc.add_paragraph()
para_space(p_footer, before=2, after=0)
r_f = p_footer.add_run('saraheid.de  ·  Website-Anleitung  ·  Mai 2026')
r_f.font.size = Pt(9)
r_f.font.color.rgb = MUTED
r_f.font.name = 'Calibri'
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ─── Speichern ───────────────────────────────────────────────

out = '/Users/michaelweiss/Desktop/Website/ANLEITUNG-STUDIO.docx'
doc.save(out)
print(f'Gespeichert: {out}')
